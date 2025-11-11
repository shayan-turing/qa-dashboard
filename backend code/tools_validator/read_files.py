import os
import pandas as pd
import io
from docx import Document

def read_excel_tools(path, sheet_name=None, selected_columns=None):
    if hasattr(path, "seek"):
        path.seek(0)

    df = pd.read_excel(path, sheet_name=sheet_name or 0)
    df.columns = [str(c).strip() for c in df.columns]

    if selected_columns:
        df = df[[col for col in selected_columns if col in df.columns]]

    print(f"✅ Loaded Excel with {len(df.columns)} columns: {df.columns.tolist()}")
    return df.to_dict(orient="records")



def read_document(input_data):
    """
    Reads and returns text from a document input.
    Handles file-like objects, bytes, paths, or plain strings.
    """

    # Case 1: It's already text
    if isinstance(input_data, str) and not os.path.exists(input_data):
        return input_data

    # Case 2: It's bytes
    if isinstance(input_data, bytes):
        return input_data.decode(errors="ignore")

    # Case 3: File-like object (Flask upload, BytesIO, etc.)
    if hasattr(input_data, "read"):
        data = input_data.read()
        try:
            return data.decode(errors="ignore")
        except Exception:
            # Possibly a DOCX binary
            input_data.seek(0)
            if input_data.name.endswith(".docx"):
                doc = Document(input_data)
                return "\n".join([p.text for p in doc.paragraphs])
            return str(data)

    # Case 4: File path
    if isinstance(input_data, str) and os.path.exists(input_data):
        if input_data.endswith(".txt") or input_data.endswith(".md"):
            with open(input_data, "r", encoding="utf-8") as f:
                return f.read()
        elif input_data.endswith(".docx"):
            doc = Document(input_data)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type")

    raise ValueError(f"Unsupported document input type: {type(input_data)}")
